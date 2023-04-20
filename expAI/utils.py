from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
import io
import matplotlib.pyplot as plt
import json
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
def gen_report_docx(trainresult, userClass="KHMT57", model="Yolov4",config= '{ "weights":"yoloface.pt", "epochs":30, "batch-size":16,"img-size": [640, 640]}	',  expName="Bài 5", expID="5234", expCreator="Đỗ Minh Hiếu",
                    expCreatedTime="2022-12-14 17:07:53", dataset="ImageNet", test_dataset_name="LFW", test_dataset_acc=0.9873, ):
    document = Document()
    black = RGBColor(0,0,0)
# parse the date string into a datetime object
    dt = datetime.strptime(str(expCreatedTime), "%Y-%m-%d %H:%M:%S.%f%z")

    # format the datetime object into the desired string format
    new_date_str = dt.strftime("%H:%M:%S %d/%m/%Y")
    heading = document.add_heading('BÁO CÁO THÍ NGHIỆM', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading('I. THÔNG TIN BÀI THÍ NGHIỆM', level=1)
    p = document.add_paragraph(f'   - Bài thí nghiệm: {expName} | Mã: {expID}')

    document.add_paragraph(f'   - Người tạo: {expCreator} - Lớp: {userClass}')

    document.add_paragraph(f'   - Thời gian tạo: {new_date_str}')

    document.add_heading('II. NỘI DUNG THỰC HÀNH', level=1)
    document.add_heading(f'  1. Mô hình lựa chọn: {model}', level=2)
    document.add_heading('  2. Cấu hình: ', level=2)
    

    # convert the dictionary to a JSON string and prettify it
    json_str = json.loads(config)
    json_str = json.dumps(json_str, indent=4)
    # add a new paragraph to the document and set the text to the prettified JSON string
    p = document.add_paragraph()
    p.add_run(json_str)
    
    document.add_heading(f'  3. Bộ dữ liệu: {dataset}', level=2)

    
    document.add_heading(
        '  4. Quá trình huấn luyện', level=2
    )
    image_stream = io.BytesIO()
    trainresultindex = [d['trainresultindex'] for d in trainresult]
    lossvalue = [d['lossvalue'] for d in trainresult]
    accuracy = [d['accuracy'] for d in trainresult]

    plt.plot(trainresultindex, lossvalue, label='Loss Value')
    plt.plot(trainresultindex, accuracy, label='Accuracy')
    plt.xlabel('Train Result Index')
    plt.ylabel('Value')
    plt.legend()
    plt.savefig(image_stream)
    image_stream.seek(0)
    p = document.add_paragraph()
    run = p.add_run()
    picture = run.add_picture(image_stream, width=Inches(5))
    p.alignment = 1
    document.add_heading(
        f'  5. Đánh giá trên tập dữ liệu: {test_dataset_name}', level=2)
    document.add_paragraph(
        f'      ✅ Độ chính xác trên tập test: {test_dataset_acc}'
    )
    
    for paragraph in document.paragraphs:
        # Loop through all the runs in the paragraph
        for run in paragraph.runs:
            # Change the color of the run to black
            run.font.color.rgb = black

    # Use XPath to select all the heading runs in the document
    heading_runs = document._element.xpath('//w:hdr//w:r') + document._element.xpath('//w:fldSimple//w:r')

    # Loop through all the heading runs in the document
    for run in heading_runs:
        # Change the color of the run to black
        run.rPr.color.rgb = black
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream
